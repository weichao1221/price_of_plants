# # -*- coding: utf-8 -*-
import os
import time

import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
import datetime
from docx.enum.section import WD_SECTION_START

import openpyxl as pl


# 封面标题（大）项目名称
def fengmian_doc1(doc, text_content: str):
    par = doc.add_paragraph("")
    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 两端对齐
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    # par.paragraph_format.first_line_indent = par.style.font.size * 2
    text = par.add_run(text_content)
    text.font.bold = False
    text.font.size = Pt(22)
    text.font.name = "Times New Roman"
    text.element.rPr.rFonts.set(qn("w:eastAsia"), u"方正小标宋简体")


# 封面标题（小）编制单位、编制时间等
def fengmian_doc2(doc, text_content: str):
    par = doc.add_paragraph("")
    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 两端对齐
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    # par.paragraph_format.first_line_indent = par.style.font.size * 2
    text = par.add_run(text_content)
    text.font.bold = False
    text.font.size = Pt(16)
    text.font.name = "Times New Roman"
    text.element.rPr.rFonts.set(qn("w:eastAsia"), u"方正小标宋简体")


# 定义一级标题格式
def Heading_1(doc, Heading_1: str):
    Heading = doc.add_heading("", level=1)
    Heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    Heading.paragraph_format.line_spacing = 1
    Heading.paragraph_format.space_before = Pt(12)
    Heading.paragraph_format.space_after = Pt(12)
    text_Heading = Heading.add_run(Heading_1)
    text_Heading.font.bold = False
    text_Heading.font.size = Pt(14)
    text_Heading.font.name = "Times New Roman"
    text_Heading.element.rPr.rFonts.set(qn("w:eastAsia"), u"黑体")
    text_Heading.font.color.rgb = RGBColor(0, 0, 0)


# 定义二级标题格式
def Heading_2(doc, Heading_2: str):
    Heading = doc.add_heading("", level=2)
    Heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    Heading.paragraph_format.line_spacing = 1
    Heading.paragraph_format.space_before = Pt(12)
    Heading.paragraph_format.space_after = Pt(12)
    # print(f"二级标题的首行缩进值{Heading.style.font.size, type(Heading.style.font.size)}")
    Heading.paragraph_format.first_line_indent = Heading.style.font.size * 2
    text_Heading = Heading.add_run(Heading_2)
    text_Heading.font.bold = True
    text_Heading.font.size = Pt(14)
    text_Heading.font.name = "Times New Roman"
    text_Heading.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")
    text_Heading.font.color.rgb = RGBColor(0, 0, 0)

# 定义三级标题格式
def Heading_3(doc, Heading_3: str):
    Heading = doc.add_heading("", level=3)
    Heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    Heading.paragraph_format.line_spacing = 1
    Heading.paragraph_format.space_before = Pt(12)
    Heading.paragraph_format.space_after = Pt(12)
    Heading.style.font.size = Pt(14)
    # print(f"三级标题的首行缩进值{Heading.style.font.size, type(Heading.style.font.size)}")
    Heading.paragraph_format.first_line_indent = Heading.style.font.size * 2
    text_Heading = Heading.add_run(Heading_3)
    text_Heading.font.bold = False
    text_Heading.font.size = Pt(14)
    text_Heading.font.name = "Times New Roman"
    text_Heading.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")
    text_Heading.font.color.rgb = RGBColor(0, 0, 0)

# 定义正文格式
def Normal_doc(doc, text_content: str):
    par = doc.add_paragraph("")
    par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 两端对齐
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    par.style.font.size = Pt(14)
    par.paragraph_format.first_line_indent = par.style.font.size * 2
    text = par.add_run(text_content)
    text.font.bold = False
    text.font.size = Pt(14)
    text.font.name = "Times New Roman"
    text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

# 结尾落款
# 结尾落款
def company_name(doc, company_name):
    par = doc.add_paragraph("")
    par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 两端对齐
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    par.paragraph_format.first_line_indent = par.style.font.size * 2
    text = par.add_run(company_name)
    text.font.bold = False
    text.font.size = Pt(14)
    text.font.name = "Times New Roman"
    text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

def created_time(doc):
    par = doc.add_paragraph("")
    par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 两端对齐
    par.paragraph_format.line_spacing = 1.5
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    par.paragraph_format.first_line_indent = par.style.font.size * 2
    text = par.add_run(f'{datetime.datetime.now().strftime("%Y年%m月%d日")}')
    text.font.bold = False
    text.font.size = Pt(14)
    text.font.name = "Times New Roman"
    text.element.rPr.rFonts.set(qn("w:eastAsia"), u"仿宋_GB2312")

# 插入新的一页
def insert_new_section(doc):
    new_section = doc.add_section(start_type=WD_SECTION_START.NEW_PAGE)
    new_section.start_type = WD_SECTION_START.EVEN_PAGE

def jdk_xd(doc, output_path, Contract_Name, Reporting_Periods, Consulting_Unit, Principal_Party, Design_Unit, Supervision_Unit, Construction_Unit, Project_Overview,
           Price_Form, Contract_Amount, Current_Application_Internal_Contract_Project, Current_Approved_Internal_Contract_Project, Current_Approved_Financial_Evaluation_Approval,
           Current_Approved_Deduction_of_Supply_of_Plants_by_Party_A, Total_Completed_Output_Value, Current_Approved_Total_Completed_Output_Value, Accounts_Payable_Advance_Payment,
           Advance_Payment_Offset, Total_Accounts_Payable, project_info_path):
    print(f'开始写{Contract_Name}的进度款审核报告···')

    # 读取附件2文档
    fj2_name = f'{Contract_Name}第{Reporting_Periods}期进度款工程付款审核表（附件二）.xlsx'
    fj2_path = os.path.join(output_path, fj2_name)
    wb_fujian2 = pl.load_workbook(fj2_path,data_only=True)
    ws_fujian2 = wb_fujian2.active

    # 读取project_info.xlsx文档
    project_info_path = os.path.join(project_info_path, 'project_info.xlsx')
    wb_project_info = pl.load_workbook(project_info_path, data_only=True)
    ws_project_info = wb_project_info.active


    for i in range(0, 4):
        Normal_doc(doc, "")
    fengmian_doc1(doc, Contract_Name)      # 此处应有变量
    fengmian_doc1(doc, f"第{Reporting_Periods}期进度款支付审核报告")  # 此处应有变量
    for i in range(0, 14):
        Normal_doc(doc, "")
    time_now = datetime.datetime.now().strftime('%Y年%m月%d日')
    fengmian_doc2(doc, f"编制单位：{Consulting_Unit}")  # 此处应有变量
    fengmian_doc2(doc, f"编制时间：{time_now}")  # 此处应有变量
    insert_new_section(doc)
    fengmian_doc2(doc, "信用承诺")  # 此处应有变量
    Normal_doc(doc, "")
    Normal_doc(doc, f"本审核报告依据相关法律法规规定、政府有关部门批复、相关技术规范、技术标准、投资政策、政府部门发布的计价标准规范、建设单"
               "位提供的项目资料编制。审核过程及审核报告文件编制坚持独立、公正、科学的原则。")
    insert_new_section(doc,)
    fengmian_doc2(doc, Contract_Name)  # 此处应有变量
    fengmian_doc2(doc, f"第{Reporting_Periods}期进度款支付审核报告")  # 此处应有变量
    Heading_1(doc, "1.概述")
    Normal_doc(doc, f"受{Principal_Party}的委托，{Consulting_Unit}对{Contract_Name}施工进行全过程造价咨询。"
               "目前开展施工进度款审核工作。在此基础上编制完成了该工程进度款审核报告。")
    Heading_1(doc, "2.项目概况")
    Heading_2(doc, "2.1 工程名称")
    Normal_doc(doc, Contract_Name)  # 此处应有变量
    Heading_2(doc, "2.2 审核性质")
    Normal_doc(doc, "施工进度款审核。")
    Heading_2(doc, "2.3 建设单位")
    Normal_doc(doc, f"{Principal_Party}")  # 此处应有变量
    Heading_2(doc, "2.4 参建单位")
    Normal_doc(doc, f"设计单位：{Design_Unit}")  # 此处应有变量
    Normal_doc(doc, f"施工单位：{Construction_Unit}")  # 此处应有变量
    Normal_doc(doc, f"监理单位：{Supervision_Unit}")  # 此处应有变量
    Heading_2(doc, "2.5 工程概况")
    Normal_doc(doc, f"{Contract_Name}。{Project_Overview}")  # 此处应有变量
    Heading_1(doc, "3.审核依据")
    Heading_2(doc, "3.1 国家有关法律、法规及定额标准文件")
    Normal_doc(doc, "（1）《建筑工程工程量清单计价规范》（GB50500-2013）；")
    Normal_doc(doc, "（2）建设工程工程量清单编制与计价规程（河北省工程建设标DB13(J)/T150-2013）；")
    Normal_doc(doc, "（3）相关费用文件，相关图集、规范等。")
    Heading_2(doc, "3.2 工程资料")
    Normal_doc(doc, "（1）招标文件、投标文件、施工合同")  #
    Normal_doc(doc, "（2）进度款申请单")
    Normal_doc(doc, f"（3）已完合同工程数量核查表（第{Reporting_Periods}期）")  # 此处应有变量
    Normal_doc(doc, "（4）补充的其他资料")
    Heading_1(doc, "4.审核说明")
    Normal_doc(doc, f"该工程为{Price_Form}，合同价款金额：{Contract_Amount}元。本次审核第{Reporting_Periods}期施工进度款，"
               "主要审核本期计量与计价金额、本期应扣减的甲供材料款、本期应付进度款、本期应扣回的"
               "预付款、本期实际支付进度款金额等内容。")
    shenjianlv = round((float(Current_Application_Internal_Contract_Project) - float(Current_Approved_Internal_Contract_Project))/ float(Current_Application_Internal_Contract_Project)* 100 ,2)
    print(f'shenjianlv:{shenjianlv}')
    Normal_doc(doc, f"4.1 本期计量与计价金额，施工单位申报金额{Current_Application_Internal_Contract_Project}元，审核金额{Current_Approved_Internal_Contract_Project}元，审减{round(float(Current_Application_Internal_Contract_Project) - float(Current_Approved_Internal_Contract_Project), 2)}元，"
               f"审减率{shenjianlv}%")
    Normal_doc(doc, f"4.2 上期财评单位扣减金额为{Current_Approved_Financial_Evaluation_Approval}元。")
    Normal_doc(doc, "4.3 本期应扣减的甲供材料款")
    Normal_doc(doc, "4.4 本期应付进度款")
    Normal_doc(doc, "依据合同第五部分专用合同条款第17.3.1付款（1）工程进度款支付金额为审核确定的当期进度款金额的85%，及第17.3.2进度款约定：（4）本期应支付比例：85%。")
    benqiyingzhifu = round((float(Current_Approved_Internal_Contract_Project) - float(Current_Approved_Financial_Evaluation_Approval) - 0) * 0.85, 2)
    Normal_doc(doc, f"本期应支付进度款=（本期计量与计价审核金额-上期财评审减-本期应扣减甲供材料款）*85%=（{Current_Approved_Internal_Contract_Project}-{Current_Approved_Financial_Evaluation_Approval}-{Current_Approved_Deduction_of_Supply_of_Plants_by_Party_A}）*85%="
               f"{benqiyingzhifu}元。")
    Normal_doc(doc, "4.5 本期应扣回的预付款")
    Normal_doc(doc, "依据合同第五部分专用合同条款第17.2.3约定预付款扣回办法：预付款在计量与计价累计金额达到签约合同价的30%之前不予扣回。在计量与计价"
               "累计金额达到签约合同价的30%之后开始扣回，在计量与计价累计金额达到签约合同价的80%前全部扣完，发包人有权根据承包履约状况，同比例从工程进度付款中扣回。")
    # 加个判断累计到本期计量与计价进度款金额占签约合同价的45.05%
    price_leiji = ws_fujian2['H15'].value
    leiji_amount = round((float(price_leiji) / float(Contract_Amount)) * 100, 2)
    if leiji_amount < 30:
        Normal_doc(doc, f"累计到本期计量与计价进度款金额占签约合同价的{leiji_amount}%，未达到预付款扣回条件。因此，本期不需要扣回预付款。")
        Current_Approved_Advance_Payment_Offset = 0
    elif 30 <= leiji_amount < 80:
        Current_Approved_Advance_Payment_Offset = round(
            (float(price_leiji) - float(Contract_Amount) * 0.3) / (float(Contract_Amount) * 0.5) * float(Accounts_Payable_Advance_Payment) - float(
                Advance_Payment_Offset), 2)
        Normal_doc(doc, f"累计到本期计量与计价进度款金额占签约合同价的{round(float(price_leiji) / float(Contract_Amount), 2) * 100}%，达到预付款扣回条件。因此，本期需要扣回预付款。")
        Normal_doc(doc, f"本期扣回预付款=（计量与计价金额累计-签约合同价*30%）/(签约合同价*80%-签约合同价*30%)*已付预付款额-上期累计扣回预付款金额"
                   f"=({price_leiji}-{Contract_Amount}*30%)/({Contract_Amount}*50%)*{Accounts_Payable_Advance_Payment}-{Advance_Payment_Offset} = {Current_Approved_Advance_Payment_Offset}元。")
        print(f"本期扣回预付款=（计量与计价金额累计-签约合同价*30%）/(签约合同价*80%-签约合同价*30%)*已付预付款额-上期累计扣回预付款金额"
                   f"=({price_leiji}-{Contract_Amount}*30%)/({Contract_Amount}*50%)*{Accounts_Payable_Advance_Payment}-{Advance_Payment_Offset} = {Current_Approved_Advance_Payment_Offset}元。")
    Normal_doc(doc, "4.6 本期实际应付进度款金额")
    benqishijizhifu = round(benqiyingzhifu - Current_Approved_Advance_Payment_Offset, 2)
    Normal_doc(doc, f"本期实际应付进度款金额=本期应支付进度款-本期扣回预付款= {benqiyingzhifu}-{Current_Approved_Advance_Payment_Offset} = {benqishijizhifu}元。")
    Normal_doc(doc, "根据合同专用条款新17.3.5条支付方式：（1）在财政资金已到位的情况下，采用银行转账方式支付。工程进度款支付采用分账支付①发包人应按照当期应"
               "付进度款的30%为农民工工资或承包人提供的人工费用数额按月单独拨付到承包人开设的农民工工资专用账户。②剩余部分支付到其他工程款结算账户。”")
    Normal_doc(doc, f"本期拨付到农民工工资专用账户金额=本期实际应付进度款金额*30%={benqishijizhifu}*30%={round(benqishijizhifu*0.3, 2)}元；")
    Normal_doc(doc, f"本期拨付到工程款结算账户金额=本期实际应付进度款金额-本期农民工工资专用账户金额={benqishijizhifu}-{round(benqishijizhifu*0.3,2)}={float(benqishijizhifu) - round(float(benqishijizhifu * 0.3), 2)}元。")
    Heading_1(doc, "5.审核结论")
    Normal_doc(doc, f"本期计量与计价金额，施工单位申报金额{Current_Application_Internal_Contract_Project}元，审核金额{Current_Approved_Internal_Contract_Project}元，其中：不含甲供苗产值{Current_Approved_Internal_Contract_Project}元，甲供苗 0元。累计审"
               f"核计量与计价金额{price_leiji}元，占合同金额{leiji_amount}%。")
    leijikouhuiyufukuan = round(float(Advance_Payment_Offset) + float(Current_Approved_Advance_Payment_Offset), 2)
    leijikouhui_advance = round(float(leijikouhuiyufukuan) / float(Accounts_Payable_Advance_Payment), 4) * 100
    Normal_doc(doc, f"本期扣回预付款{Current_Approved_Advance_Payment_Offset}元，累计本期扣回预付款{leijikouhuiyufukuan}元，占预付款总额{leijikouhui_advance}%。")
    Normal_doc(doc, f"上一期财评单位扣减{Current_Approved_Financial_Evaluation_Approval}元。")
    leijishijiyingzhifu = round(float(Total_Accounts_Payable) + float(benqishijizhifu), 2)
    leijishijiyingzhifu_amount = round(float(leijishijiyingzhifu) / float(Contract_Amount), 4) * 100
    Normal_doc(doc, f"本期实际应付进度款金额{benqishijizhifu}元，累计实际应付金额{leijishijiyingzhifu}元，占合同金额{leijishijiyingzhifu_amount}%。")
    Normal_doc(doc, f"本期实际应付进度款，拨到农民工工资专用账户金额{round(benqishijizhifu*0.3, 2)}元，拨到工程款结算账户金额{float(benqishijizhifu) - round(float(benqishijizhifu * 0.3), 2)}元。")
    Heading_1(doc, "6.其他说明")
    Normal_doc(doc, "此审核金额只用于进度款的支付，不作为结算依据。")
    Normal_doc(doc, "")
    Normal_doc(doc, "")
    Normal_doc(doc, "")
    company_name(doc, Consulting_Unit)    # 此处应有变量
    created_time(doc, )
    file_name = f"{Contract_Name}第{Reporting_Periods}期进度款支付审核报告{time_now}.docx"
    save_path = os.path.join(output_path, file_name)
    return save_path


